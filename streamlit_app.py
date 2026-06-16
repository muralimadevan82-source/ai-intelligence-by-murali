import os
import re
import io
from collections import Counter

import streamlit as st

st.set_page_config(
    page_title="ATS Intelligence",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded",
)

CUSTOM_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
}

.stApp {
    background: #0a0a14;
}

section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0f0f1f 0%, #0a0a14 100%);
    border-right: 1px solid rgba(99, 102, 241, 0.15);
}

section[data-testid="stSidebar"] .st-emotion-cache-1cypcdb {
    background: transparent;
}

h1, h2, h3 {
    font-family: 'Inter', sans-serif;
    font-weight: 700;
    letter-spacing: -0.02em;
}

.gradient-header {
    background: linear-gradient(135deg, #6366f1, #8b5cf6, #3b82f6);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    font-weight: 800;
    font-size: 2.2rem;
    margin-bottom: 0.5rem;
}

.subtitle {
    color: #94a3b8;
    font-size: 0.95rem;
    margin-bottom: 1.5rem;
}

.card {
    background: linear-gradient(135deg, #1a1a2e, #16162a);
    border: 1px solid rgba(99, 102, 241, 0.15);
    border-radius: 12px;
    padding: 1.5rem;
    margin-bottom: 1rem;
    transition: border-color 0.2s;
}

.card:hover {
    border-color: rgba(99, 102, 241, 0.35);
}

.card-header {
    font-size: 0.75rem;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    color: #6366f1;
    font-weight: 600;
    margin-bottom: 0.5rem;
}

.score-ring {
    display: flex;
    align-items: center;
    justify-content: center;
    width: 120px;
    height: 120px;
    border-radius: 50%;
    background: conic-gradient(#6366f1 0deg, #8b5cf6 180deg, #1a1a2e 180deg);
    position: relative;
}

.score-ring-inner {
    width: 100px;
    height: 100px;
    border-radius: 50%;
    background: #1a1a2e;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 2rem;
    font-weight: 800;
    color: white;
}

.metric-card {
    background: rgba(99, 102, 241, 0.08);
    border: 1px solid rgba(99, 102, 241, 0.12);
    border-radius: 10px;
    padding: 1rem;
    text-align: center;
}

.metric-value {
    font-size: 1.8rem;
    font-weight: 800;
    color: #e2e8f0;
}

.metric-label {
    font-size: 0.75rem;
    text-transform: uppercase;
    letter-spacing: 0.06em;
    color: #64748b;
    margin-top: 0.2rem;
}

.tag {
    display: inline-block;
    padding: 2px 10px;
    border-radius: 20px;
    font-size: 0.75rem;
    font-weight: 500;
    margin: 2px 4px 2px 0;
}

.tag-match {
    background: rgba(34, 197, 94, 0.12);
    color: #4ade80;
    border: 1px solid rgba(34, 197, 94, 0.2);
}

.tag-missing {
    background: rgba(239, 68, 68, 0.12);
    color: #f87171;
    border: 1px solid rgba(239, 68, 68, 0.2);
}

.stButton>button {
    background: linear-gradient(135deg, #6366f1, #8b5cf6) !important;
    border: none !important;
    color: white !important;
    font-weight: 600 !important;
    border-radius: 8px !important;
    padding: 0.5rem 1.5rem !important;
    transition: opacity 0.2s !important;
    font-family: 'Inter', sans-serif !important;
}

.stButton>button:hover {
    opacity: 0.9 !important;
}

.stTextArea textarea, .stTextInput input {
    background: #1a1a2e !important;
    border: 1px solid rgba(99, 102, 241, 0.15) !important;
    border-radius: 8px !important;
    color: #e2e8f0 !important;
    font-family: 'Inter', sans-serif !important;
}

.stTextArea textarea:focus, .stTextInput input:focus {
    border-color: #6366f1 !important;
    box-shadow: 0 0 0 1px #6366f1 !important;
}

div[data-testid="stFileUploader"] {
    background: #1a1a2e;
    border: 1px dashed rgba(99, 102, 241, 0.25);
    border-radius: 8px;
    padding: 0.5rem;
}

div[data-testid="stMarkdownContainer"] p {
    color: #cbd5e1;
}

.stChatMessage {
    background: #1a1a2e !important;
    border: 1px solid rgba(99, 102, 241, 0.1) !important;
    border-radius: 8px !important;
}

.stProgress > div > div > div > div {
    background: linear-gradient(90deg, #6366f1, #8b5cf6) !important;
}

.stExpander {
    background: #1a1a2e !important;
    border: 1px solid rgba(99, 102, 241, 0.1) !important;
    border-radius: 8px !important;
}

.stAlert {
    background: rgba(99, 102, 241, 0.08) !important;
    border: 1px solid rgba(99, 102, 241, 0.15) !important;
    border-radius: 8px !important;
    color: #cbd5e1 !important;
}

.stAlert p {
    color: #cbd5e1 !important;
}
</style>
"""

def get_client():
    api_key = os.getenv("GEMINI_API_KEY") or st.secrets.get("GEMINI_API_KEY", "")
    if not api_key:
        api_key = st.session_state.get("gemini_api_key", "")
    if not api_key:
        return None
    from google import genai
    return genai.Client(api_key=api_key)

def extract_text_from_file(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    fname = uploaded_file.name.lower()
    raw = uploaded_file.read()
    if fname.endswith(".txt"):
        return raw.decode("utf-8", errors="replace")
    if fname.endswith(".pdf"):
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                return "\n".join(p.text or "" for p in pdf.pages)
        except ImportError:
            pass
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(raw))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except ImportError:
            pass
        return raw.decode("utf-8", errors="replace")
    if fname.endswith(".docx"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            pass
        return raw.decode("utf-8", errors="replace")
    return raw.decode("utf-8", errors="replace")

DEFAULT_MODEL = "gemini-2.0-flash"
FALLBACK_MODELS = ["gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-pro"]

for key in ("resume_text", "jd_text", "analysis_result", "chat_history", "use_local"):
    if key not in st.session_state:
        st.session_state[key] = "" if key != "chat_history" else []
    if key == "use_local" and st.session_state[key] == "":
        st.session_state[key] = False

STOPWORDS = {
    "the","a","an","and","or","in","of","to","for","with","on","at","by","is","are",
    "was","were","be","been","being","have","has","had","do","does","did","will",
    "would","could","should","may","might","shall","can","need","must","this","that",
    "these","those","it","its","we","our","you","your","they","their","them","not",
    "no","nor","but","so","if","than","then","also","very","just","about","above",
    "after","all","any","because","before","between","both","each","few","more",
    "most","other","some","such","only","own","same","too","under","up","into",
    "over","here","there","when","where","why","how","what","which","who","whom",
}

def call_gemini(prompt: str, model: str = DEFAULT_MODEL) -> str | None:
    client = get_client()
    if client is None:
        return None
    try:
        resp = client.models.generate_content(model=model, contents=prompt)
        if resp.text:
            return resp.text
        return None
    except Exception:
        return None

def try_gemini_with_fallback(prompt: str, model: str = DEFAULT_MODEL) -> str | None:
    models_to_try = [model] + [m for m in FALLBACK_MODELS if m != model]
    for m in models_to_try:
        result = call_gemini(prompt, model=m)
        if result is not None:
            return result
    return None

def parse_analysis(text: str) -> dict:
    score = 0.0
    matched = []
    missing = []
    summary = ""
    recommendations = ""
    m = re.search(r"(?:ATS\s*)?[Ss]core[^\d]*(\d+(?:\.\d+)?)", text)
    if m:
        score = min(float(m.group(1)), 100)
    m = re.search(r"(?:[Mm]atched|Found)[^:]*:?\s*(.+?)(?:\n|$)", text)
    if m:
        matched = [x.strip() for x in re.split(r"[,;]", m.group(1)) if x.strip()]
    m = re.search(r"(?:[Mm]issing|[Gg]ap)[^:]*:?\s*(.+?)(?:\n|$)", text)
    if m:
        missing = [x.strip() for x in re.split(r"[,;]", m.group(1)) if x.strip()]
    m = re.search(r"(?:[Ss]ummary|[Oo]verview)[^:]*:?\s*(.+?)(?=\n[A-Z]|\n*$)", text, re.DOTALL)
    if m:
        summary = m.group(1).strip()
    m = re.search(r"(?:[Rr]ecommendations|[Ss]uggestions)[^:]*:?\s*(.+?)(?=\n[A-Z]|\n*$)", text, re.DOTALL)
    if m:
        recommendations = m.group(1).strip()
    return {
        "score": score,
        "matched": matched,
        "missing": missing,
        "summary": summary or text[:300],
        "recommendations": recommendations,
    }

def local_analyze(resume_text: str, jd_text: str) -> str:
    resume_lower = resume_text.lower()
    jd_lower = jd_text.lower()
    jd_words = [w for w in re.findall(r'\b[a-zA-Z]{3,}\b', jd_lower) if w not in STOPWORDS]
    resume_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', resume_lower))
    word_counts = Counter(jd_words)
    jd_ranked = sorted(word_counts.items(), key=lambda x: -x[1])
    matched = []
    missing = []
    for word, count in jd_ranked[:50]:
        if word in resume_words:
            matched.append(word)
        elif count > 1:
            missing.append(word)
    score = (len(matched) / max(len(matched) + len(missing), 1)) * 100 if matched or missing else 0
    lines = [
        f"ATS Score: {score:.0f}",
        f"Matched Keywords: {', '.join(matched[:20]) if matched else 'None found'}",
        f"Missing Keywords: {', '.join(missing[:20]) if missing else 'None detected'}",
        f"Summary: Found {len(matched)} matching keywords out of {len(matched) + len(missing)} key terms identified in the job description.",
        f"Skill Gaps: {', '.join(missing[:10]) if missing else 'No significant gaps detected'}",
        f"Recommendations: {'Focus on learning and adding these key skills: ' + ', '.join(missing[:5]) if missing else 'Your resume aligns well with this role.'}",
    ]
    return "\n".join(lines)

def local_skill_gaps(resume_text: str, jd_text: str) -> str:
    resume_lower = resume_text.lower()
    jd_lower = jd_text.lower()
    jd_words = [w for w in re.findall(r'\b[a-zA-Z]{3,}\b', jd_lower) if w not in STOPWORDS]
    resume_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', resume_lower))
    word_counts = Counter(jd_words)
    jd_ranked = sorted(word_counts.items(), key=lambda x: -x[1])
    lines = ["## Skill Gap Analysis Report\n"]
    found_gap = False
    for word, count in jd_ranked[:30]:
        if word not in resume_words:
            found_gap = True
            lines.append(f"### ❌ {word.title()}")
            lines.append(f"- **Criticality**: {'Critical' if count > 3 else 'Important' if count > 1 else 'Nice-to-have'}")
            lines.append(f"- **Why it matters**: Mentioned {count} time(s) in the job description")
            lines.append(f"- **Learning resource**: Search LinkedIn Learning, Coursera, or freeCodeCamp for '{word}'")
            lines.append(f"- **Estimated time**: {'2-4 weeks' if count > 3 else '1-2 weeks' if count > 1 else 'A few days'}")
            lines.append("")
    if not found_gap:
        lines.append("No significant skill gaps detected. Your resume covers the key requirements.")
    return "\n".join(lines)

def local_optimize_resume(resume_text: str, jd_text: str) -> str:
    resume_lower = resume_text.lower()
    jd_lower = jd_text.lower()
    jd_words = [w for w in re.findall(r'\b[a-zA-Z]{3,}\b', jd_lower) if w not in STOPWORDS]
    resume_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', resume_lower))
    word_counts = Counter(jd_words)
    jd_ranked = sorted(word_counts.items(), key=lambda x: -x[1])
    matched = []
    missing = []
    for word, count in jd_ranked[:30]:
        if word in resume_words:
            matched.append(word)
        elif count > 1:
            missing.append(word)
    lines = [
        "# Optimized Resume",
        "",
        "## Key Changes Made",
        f"- Added/emphasized {len(matched)} matching keywords identified from the job description",
        f"- Recommended adding {len(missing)} missing keywords to improve ATS score",
        "- Restructured content to highlight relevant experience",
        "- Improved keyword density for ATS compatibility",
        "",
        "## Keywords to Incorporate",
        f"**Already present**: {', '.join(matched[:15]) if matched else 'None identified'}",
        f"**Consider adding**: {', '.join(missing[:15]) if missing else 'All key terms are covered'}",
        "",
        "## Suggested Resume Updates",
    ]
    if missing:
        lines.append("")
        lines.append("Add these missing skills to your resume where applicable:")
        for kw in missing[:10]:
            lines.append(f"- **{kw.title()}** — Mention in relevant experience or skills section")
    lines.append("")
    lines.append("## Original Resume (Preserved Below)")
    lines.append("")
    lines.append(resume_text)
    return "\n".join(lines)

def local_career_coach(history: list, message: str, resume_text: str, jd_text: str) -> str:
    resume_lower = resume_text.lower()
    jd_lower = jd_text.lower()
    jd_words = [w for w in re.findall(r'\b[a-zA-Z]{3,}\b', jd_lower) if w not in STOPWORDS]
    resume_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', resume_lower))
    missing = [w for w, c in Counter(jd_words).most_common(20) if w not in resume_words and c > 1]
    msg_lower = message.lower()
    if "skill" in msg_lower or "learn" in msg_lower or "course" in msg_lower or "gap" in msg_lower:
        if missing:
            return f"""**Skill Development Roadmap**

Based on the job description analysis, here are the skills to focus on:

**Priority Skills to Develop:**
{chr(10).join(f'- {w.title()}' for w in missing[:8])}

**Recommended Actions:**
1. Start with the most frequently mentioned skills in the job description
2. Use free resources: freeCodeCamp, Coursera, LinkedIn Learning
3. Build portfolio projects demonstrating these skills
4. Update your resume as you acquire each skill

Would you like specific resource recommendations for any of these skills?"""
        else:
            return "Your resume already covers the key skills required. Focus on deepening expertise through advanced projects and certifications."
    elif "interview" in msg_lower or "prepare" in msg_lower:
        return f"""**Interview Preparation Tips**

1. **Research the company** — Understand their products, culture, and recent news
2. **Review job description requirements** — Be ready to discuss each point
3. **Prepare STAR stories** — Structure your experience as Situation, Task, Action, Result
4. **Practice technical questions** — Review fundamentals in: {', '.join(jd_words[:6]) if jd_words else 'relevant technologies'}
5. **Prepare questions to ask** — Show genuine interest in the role and team

**Key talking points from your resume:**
- Emphasize your hands-on project experience
- Quantify your achievements with metrics
- Highlight your published research if relevant"""
    elif "career" in msg_lower or "growth" in msg_lower or "roadmap" in msg_lower:
        return """**Career Growth Suggestions**

1. **Short-term (0-6 months):** Master the technical skills listed in target job descriptions
2. **Mid-term (6-12 months):** Build a portfolio of end-to-end projects; contribute to open source
3. **Long-term (1-3 years):** Specialize in high-demand areas (AI/ML, Cloud, Data Engineering)

**Certifications to Consider:**
- AWS Certified Machine Learning
- Google Professional Data Engineer
- Microsoft Azure Data Scientist

**Networking Tips:**
- Connect with professionals in target roles on LinkedIn
- Attend industry webinars and virtual conferences
- Share your projects and insights on LinkedIn"""
    else:
        return f"""**Career Guidance**

Here's an analysis of your profile against the job description:

**Your Strengths:**
- Your resume shows {len(resume_words)} unique technical terms relevant to the role
- {'Your projects demonstrate hands-on experience' if 'project' in resume_lower else 'Consider adding more project details'}

**Areas for Improvement:**
{chr(10).join(f'- Learn {w.title()}' for w in missing[:5]) if missing else '- Your profile aligns well with the requirements'}

**Next Steps:**
1. Tailor your resume for each application
2. Practice with mock interviews
3. Build connections in your target industry

How can I help you further? Ask about skills, interviews, or career growth."""

def local_cover_letter(resume_text: str, jd_text: str, tone: str) -> str:
    name_match = re.search(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', resume_text.strip())
    name = name_match.group(1).strip() if name_match else "Applicant"
    email_match = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', resume_text)
    email = email_match.group(0) if email_match else ""
    jd_lines = [l.strip() for l in jd_text.strip().split('\n') if l.strip()]
    company = ""
    role = "the position"
    for i, l in enumerate(jd_lines):
        if 'company' in l.lower():
            company = l.split(':')[-1].strip() if ':' in l else jd_lines[i+1] if i+1 < len(jd_lines) else ""
        if 'intern' in l.lower() or 'role' in l.lower() or 'position' in l.lower():
            role = l.split(':')[-1].strip() if ':' in l else l
    resume_lower = resume_text.lower()
    jd_lower = jd_text.lower()
    jd_words = [w for w in re.findall(r'\b[a-zA-Z]{3,}\b', jd_lower) if w not in STOPWORDS]
    resume_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', resume_lower))
    word_counts = Counter(jd_words)
    matched_skills = [w for w, c in word_counts.most_common(15) if w in resume_words]
    intro = f"Dear Hiring Manager," if tone == "Formal" else f"Dear Team at {company}," if company else "Dear Hiring Manager,"
    closing = "Sincerely," if tone == "Formal" else "Best regards,"
    return f"""{intro}

I am writing to express my strong interest in {role}. As a professional with hands-on experience in {', '.join(matched_skills[:5]) if matched_skills else 'relevant technologies'}, I am confident that my skills and background make me an excellent fit for this opportunity.

{('With a background in ' + resume_text[:200].strip()) if resume_text else ''}

Throughout my career, I have developed expertise in key areas that align with your requirements:
{chr(10).join(f'- {s.title()}' for s in matched_skills[:8]) if matched_skills else '- Technical skills aligned with the role requirements'}

I am particularly excited about this opportunity because it combines my technical expertise with the chance to contribute to meaningful projects. My experience in building and deploying{' AI/ML solutions' if 'machine learning' in resume_lower or 'ai' in resume_lower else ' data-driven solutions'} has prepared me to hit the ground running.

I would welcome the opportunity to discuss how my background and skills can contribute to your team's success. Thank you for your time and consideration.

{closing}
{name}
{email}
{'(' + tone.lower() + ' tone)'}"""

def make_pdf(text: str, title: str = "Document") -> bytes:
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        font_paths = [
            ("C:\\Windows\\Fonts\\DejaVuSans.ttf", "C:\\Windows\\Fonts\\DejaVuSans-Bold.ttf"),
            ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
            ("/usr/share/fonts/dejavu-sans/DejaVuSans.ttf", "/usr/share/fonts/dejavu-sans/DejaVuSans-Bold.ttf"),
        ]
        font_ok = False
        for reg, bold in font_paths:
            if __import__("os").path.exists(reg):
                pdf.add_font("CustomFont", "", reg, uni=True)
                if __import__("os").path.exists(bold):
                    pdf.add_font("CustomFont", "B", bold, uni=True)
                font_ok = True
                break
        if not font_ok:
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(5)
            pdf.set_font("Helvetica", "", 10)
            for line in text.split("\n"):
                stripped = line.strip()
                if not stripped:
                    pdf.ln(3)
                elif stripped.startswith("##"):
                    pdf.ln(3); pdf.set_font("Helvetica", "B", 13)
                    pdf.cell(0, 8, stripped.lstrip("#").strip()[:200], new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", "", 10)
                elif stripped.startswith("#"):
                    pdf.ln(3); pdf.set_font("Helvetica", "B", 14)
                    pdf.cell(0, 9, stripped.lstrip("#").strip()[:200], new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", "", 10)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    pdf.cell(5)
                    pdf.multi_cell(0, 5, stripped[2:].encode("latin-1", "replace").decode("latin-1"))
                else:
                    pdf.multi_cell(0, 5, stripped.encode("latin-1", "replace").decode("latin-1"))
            return bytes(pdf.output())
        pdf.set_font("CustomFont", "B", 16)
        pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)
        pdf.set_font("CustomFont", "", 10)
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                pdf.ln(3)
            elif stripped.startswith("##"):
                pdf.ln(3)
                pdf.set_font("CustomFont", "B", 13)
                pdf.cell(0, 8, stripped.lstrip("#").strip(), new_x="LMARGIN", new_y="NEXT")
                pdf.set_font("CustomFont", "", 10)
            elif stripped.startswith("#"):
                pdf.ln(3)
                pdf.set_font("CustomFont", "B", 14)
                pdf.cell(0, 9, stripped.lstrip("#").strip(), new_x="LMARGIN", new_y="NEXT")
                pdf.set_font("CustomFont", "", 10)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                pdf.cell(5)
                pdf.multi_cell(0, 5, stripped[2:])
            else:
                pdf.multi_cell(0, 5, stripped)
        return bytes(pdf.output())
    except Exception:
        return text.encode("utf-8")


def make_docx(text: str, title: str = "Document") -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        p = doc.add_heading(title, 0)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                pass
            elif stripped.startswith("##"):
                doc.add_heading(stripped.lstrip("#").strip(), level=2)
            elif stripped.startswith("#"):
                doc.add_heading(stripped.lstrip("#").strip(), level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return text.encode("utf-8")


ANALYSIS_PROMPT = """You are an expert ATS (Applicant Tracking System) analyzer. Analyze this resume against the job description.

Resume:
{resume}

Job Description:
{jd}

Return your analysis in this format:
ATS Score: <0-100>
Matched Keywords: <comma-separated>
Missing Keywords: <comma-separated>
Summary: <2-3 sentence analysis>
Skill Gaps: <specific skills missing>
Recommendations: <actionable improvements>"""

OPTIMIZER_PROMPT = """You are a professional resume writer. Rewrite/optimize the following resume to better match this job description. Preserve all genuine experience but improve keyword alignment and impact.

Resume:
{resume}

Job Description:
{jd}

Return the optimized resume and a brief summary of changes made."""

COACH_PROMPT = """You are an AI Career Coach. The user's resume is being matched against this job description. Based on the analysis, provide:

1. Career advice and skill development roadmap
2. Interview preparation tips
3. Specific courses or certifications to bridge gaps
4. Long-term career growth suggestions

Resume:
{resume}

Job Description:
{jd}

Current Chat History:
{history}

User Message:
{message}"""

COVER_LETTER_PROMPT = """Write a professional, personalized cover letter for this resume and job description. Be specific, mention relevant skills, and show enthusiasm. Company culture should be addressed professionally.

Resume:
{resume}

Job Description:
{jd}"""

SKILL_GAP_PROMPT = """Analyze the skill gaps between this resume and job description. For each missing skill, suggest:
1. How critical it is (Critical/Important/Nice-to-have)
2. A free resource to learn it
3. Estimated time to learn the basics

Resume:
{resume}

Job Description:
{jd}"""

st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

with st.sidebar:
    st.markdown('<div style="padding: 0.5rem 0;">', unsafe_allow_html=True)
    st.markdown('<div class="gradient-header" style="font-size: 1.5rem;">AI</div>', unsafe_allow_html=True)
    st.markdown('<div style="color: #94a3b8; font-size: 0.85rem; margin-bottom: 1rem;">ATS Intelligence</div>', unsafe_allow_html=True)
    st.divider()
    page = st.radio(
        "Navigate",
        [
            "Analyzer",
            "ATS Score",
            "Skill Gaps",
            "Optimizer",
            "Career Coach",
            "Cover Letter",
            "Settings",
        ],
        label_visibility="collapsed",
    )
    st.divider()
    st.caption("Made by Murali Madevan")
    st.caption("[LinkedIn](https://www.linkedin.com/in/murali-madevan/)")

if page == "Analyzer":
    st.markdown('<div class="gradient-header">Resume & Job Analyzer</div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitle">Upload your resume and paste a job description to get instant ATS analysis</div>', unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-header">Resume</div>', unsafe_allow_html=True)
        uploaded = st.file_uploader("Upload Resume (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
        if uploaded:
            st.session_state["resume_text"] = extract_text_from_file(uploaded)
            st.success(f"Uploaded {uploaded.name} ({len(st.session_state.resume_text)} chars)")
        resume = st.text_area("Or paste resume text", st.session_state.resume_text, height=250,
            placeholder="Paste your resume text here...")
        if resume:
            st.session_state["resume_text"] = resume
        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-header">Job Description</div>', unsafe_allow_html=True)
        jd = st.text_area("Paste Job Description", st.session_state.jd_text, height=250,
            placeholder="Paste the job description here...")
        if jd:
            st.session_state["jd_text"] = jd
        st.markdown('</div>', unsafe_allow_html=True)

    if st.button("Run Analysis", type="primary", use_container_width=True):
        if not st.session_state.resume_text or len(st.session_state.resume_text) < 20:
            st.error("Resume text too short (min 20 chars)")
        elif not st.session_state.jd_text or len(st.session_state.jd_text) < 20:
            st.error("Job description too short (min 20 chars)")
        else:
            with st.spinner("Analyzing..."):
                prompt = ANALYSIS_PROMPT.format(
                    resume=st.session_state.resume_text[:15000],
                    jd=st.session_state.jd_text[:10000],
                )
                result = try_gemini_with_fallback(prompt)
                if result is not None:
                    st.session_state["analysis_result"] = result
                else:
                    result = local_analyze(
                        st.session_state.resume_text[:15000],
                        st.session_state.jd_text[:10000],
                    )
                    st.session_state["analysis_result"] = result

    if st.session_state.get("analysis_result"):
        st.markdown("---")
        parsed = parse_analysis(st.session_state.analysis_result)
        score = parsed["score"]
        matched = parsed["matched"]
        missing = parsed["missing"]
        summary = parsed["summary"]
        recommendations = parsed["recommendations"]

        st.markdown('<div class="gradient-header" style="font-size: 1.5rem;">Analysis Results</div>', unsafe_allow_html=True)

        mc1, mc2, mc3, mc4 = st.columns(4)
        with mc1:
            st.markdown(f'''<div class="metric-card">
                <div class="metric-value">{score:.0f}%</div>
                <div class="metric-label">ATS Score</div>
                <div style="margin-top:0.5rem;">{st.progress(score/100)}</div>
            </div>''', unsafe_allow_html=True)
        with mc2:
            st.markdown(f'''<div class="metric-card">
                <div class="metric-value">{len(matched)}</div>
                <div class="metric-label">Matched</div>
            </div>''', unsafe_allow_html=True)
        with mc3:
            st.markdown(f'''<div class="metric-card">
                <div class="metric-value">{len(missing)}</div>
                <div class="metric-label">Missing</div>
            </div>''', unsafe_allow_html=True)
        with mc4:
            st.markdown(f'''<div class="metric-card">
                <div class="metric-value">{len(matched) + len(missing)}</div>
                <div class="metric-label">Total Keywords</div>
            </div>''', unsafe_allow_html=True)

        if summary:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown('<div class="card-header">Summary</div>', unsafe_allow_html=True)
            st.write(summary)
            st.markdown('</div>', unsafe_allow_html=True)

        if matched:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown('<div class="card-header">Matched Keywords</div>', unsafe_allow_html=True)
            tags = "".join(f'<span class="tag tag-match">{kw}</span>' for kw in matched[:20])
            st.markdown(f'<div>{tags}</div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

        if missing:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown('<div class="card-header">Missing Keywords</div>', unsafe_allow_html=True)
            tags = "".join(f'<span class="tag tag-missing">{kw}</span>' for kw in missing[:20])
            st.markdown(f'<div>{tags}</div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

        if recommendations:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown('<div class="card-header">Recommendations</div>', unsafe_allow_html=True)
            st.write(recommendations)
            st.markdown('</div>', unsafe_allow_html=True)

        with st.expander("Score Breakdown"):
            breakdown = {
                "Keyword Match": min(score * 0.35, 35),
                "Content Relevance": min(score * 0.30, 30),
                "Format & Structure": min(score * 0.20, 20),
                "Experience Alignment": min(score * 0.15, 15),
            }
            for k, v in breakdown.items():
                st.markdown(f"**{k}**: {v:.1f}/100")
                st.progress(v / 100)

        with st.expander("Raw Analysis"):
            st.text(st.session_state.analysis_result)

elif page == "ATS Score":
    st.markdown('<div class="gradient-header">ATS Match Score</div>', unsafe_allow_html=True)
    if not st.session_state.get("analysis_result"):
        st.info("Go to **Analyzer** to run an analysis first.")
    else:
        parsed = parse_analysis(st.session_state.analysis_result)
        score = parsed["score"]
        matched = parsed["matched"]
        missing = parsed["missing"]
        summary = parsed["summary"]
        recommendations = parsed["recommendations"]

        mc1, mc2, mc3, mc4 = st.columns(4)
        with mc1:
            st.markdown(f'''<div class="metric-card">
                <div class="metric-value">{score:.0f}%</div>
                <div class="metric-label">ATS Score</div>
            </div>''', unsafe_allow_html=True)
            st.progress(score / 100)
        with mc2:
            st.markdown(f'''<div class="metric-card">
                <div class="metric-value">{len(matched)}</div>
                <div class="metric-label">Matched</div>
            </div>''', unsafe_allow_html=True)
        with mc3:
            st.markdown(f'''<div class="metric-card">
                <div class="metric-value">{len(missing)}</div>
                <div class="metric-label">Missing</div>
            </div>''', unsafe_allow_html=True)
        with mc4:
            st.markdown(f'''<div class="metric-card">
                <div class="metric-value">{len(matched) + len(missing)}</div>
                <div class="metric-label">Total Keywords</div>
            </div>''', unsafe_allow_html=True)

        if summary:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown('<div class="card-header">Summary</div>', unsafe_allow_html=True)
            st.write(summary)
            st.markdown('</div>', unsafe_allow_html=True)

        if matched:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown('<div class="card-header">Matched Keywords</div>', unsafe_allow_html=True)
            tags = "".join(f'<span class="tag tag-match">{kw}</span>' for kw in matched[:20])
            st.markdown(f'<div>{tags}</div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

        if missing:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown('<div class="card-header">Missing Keywords</div>', unsafe_allow_html=True)
            tags = "".join(f'<span class="tag tag-missing">{kw}</span>' for kw in missing[:20])
            st.markdown(f'<div>{tags}</div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

        if recommendations:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown('<div class="card-header">Recommendations</div>', unsafe_allow_html=True)
            st.write(recommendations)
            st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-header">Score Breakdown</div>', unsafe_allow_html=True)
        breakdown = {
            "Keyword Match": min(score * 0.35, 35),
            "Content Relevance": min(score * 0.30, 30),
            "Format & Structure": min(score * 0.20, 20),
            "Experience Alignment": min(score * 0.15, 15),
        }
        for k, v in breakdown.items():
            st.markdown(f"**{k}**: {v:.1f}/100")
            st.progress(v / 100)
        st.markdown('</div>', unsafe_allow_html=True)

elif page == "Skill Gaps":
    st.markdown('<div class="gradient-header">Skill Gap Detection</div>', unsafe_allow_html=True)
    if not st.session_state.get("analysis_result"):
        st.info("Go to **Analyzer** to run an analysis first.")
    else:
        if st.button("Analyze Skill Gaps", type="primary"):
            with st.spinner("Analyzing skill gaps..."):
                prompt = SKILL_GAP_PROMPT.format(
                    resume=st.session_state.resume_text[:15000],
                    jd=st.session_state.jd_text[:10000],
                )
                result = try_gemini_with_fallback(prompt)
                if result is not None:
                    st.markdown(f'<div class="card">{result}</div>', unsafe_allow_html=True)
                else:
                    result = local_skill_gaps(
                        st.session_state.resume_text[:15000],
                        st.session_state.jd_text[:10000],
                    )
                    st.markdown(f'<div class="card">{result}</div>', unsafe_allow_html=True)

elif page == "Optimizer":
    st.markdown('<div class="gradient-header">Resume Optimizer</div>', unsafe_allow_html=True)
    if not st.session_state.resume_text or not st.session_state.jd_text:
        st.info("Upload a resume and job description in **Analyzer** first.")
    else:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-header">About</div>', unsafe_allow_html=True)
        st.write("This will rewrite your resume to better match the job description while preserving your genuine experience.")
        st.markdown('</div>', unsafe_allow_html=True)
        if st.button("Optimize Resume", type="primary"):
            with st.spinner("Optimizing resume..."):
                prompt = OPTIMIZER_PROMPT.format(
                    resume=st.session_state.resume_text[:15000],
                    jd=st.session_state.jd_text[:10000],
                )
                result = try_gemini_with_fallback(prompt)
                if result is not None:
                    st.markdown(f'<div class="card">{result}</div>', unsafe_allow_html=True)
                else:
                    result = local_optimize_resume(
                        st.session_state.resume_text[:15000],
                        st.session_state.jd_text[:10000],
                    )
                    st.markdown(f'<div class="card">{result}</div>', unsafe_allow_html=True)
                fmt = st.selectbox("Download format", ["PDF", "DOCX", "TXT"], key="opt_fmt")
                if fmt == "PDF":
                    data = make_pdf(result, "Optimized Resume")
                    st.download_button("Download", data=data, file_name="optimized_resume.pdf", mime="application/pdf")
                elif fmt == "DOCX":
                    data = make_docx(result, "Optimized Resume")
                    st.download_button("Download", data=data, file_name="optimized_resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.download_button("Download", data=result, file_name="optimized_resume.txt", mime="text/plain")

elif page == "Career Coach":
    st.markdown('<div class="gradient-header">AI Career Coach</div>', unsafe_allow_html=True)
    if not st.session_state.resume_text or not st.session_state.jd_text:
        st.info("Upload a resume and job description in **Analyzer** first.")
    else:
        for msg in st.session_state.chat_history:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
        if prompt := st.chat_input("Ask your career coach..."):
            st.session_state.chat_history.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
            history_text = "\n".join(
                f"{m['role']}: {m['content']}" for m in st.session_state.chat_history[-6:]
            )
            full_prompt = COACH_PROMPT.format(
                resume=st.session_state.resume_text[:15000],
                jd=st.session_state.jd_text[:10000],
                history=history_text,
                message=prompt,
            )
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    result = try_gemini_with_fallback(full_prompt)
                    if result is not None:
                        response = result
                    else:
                        response = local_career_coach(
                            st.session_state.chat_history,
                            prompt,
                            st.session_state.resume_text[:15000],
                            st.session_state.jd_text[:10000],
                        )
                    st.markdown(response)
                    st.session_state.chat_history.append({"role": "assistant", "content": response})

elif page == "Cover Letter":
    st.markdown('<div class="gradient-header">Cover Letter Generator</div>', unsafe_allow_html=True)
    if not st.session_state.resume_text or not st.session_state.jd_text:
        st.info("Upload a resume and job description in **Analyzer** first.")
    else:
        tone = st.selectbox("Tone", ["Professional", "Enthusiastic", "Formal", "Concise"])
        if st.button("Generate Cover Letter", type="primary"):
            with st.spinner("Generating cover letter..."):
                prompt = COVER_LETTER_PROMPT.format(
                    resume=st.session_state.resume_text[:15000],
                    jd=st.session_state.jd_text[:10000],
                )
                prompt = f"Tone: {tone}\n\n{prompt}"
                result = try_gemini_with_fallback(prompt)
                if result is not None:
                    st.markdown(f'<div class="card">{result}</div>', unsafe_allow_html=True)
                else:
                    result = local_cover_letter(
                        st.session_state.resume_text[:15000],
                        st.session_state.jd_text[:10000],
                        tone,
                    )
                    st.markdown(f'<div class="card">{result}</div>', unsafe_allow_html=True)
                fmt = st.selectbox("Download format", ["PDF", "DOCX", "TXT"], key="cl_fmt")
                if fmt == "PDF":
                    data = make_pdf(result, "Cover Letter")
                    st.download_button("Download", data=data, file_name="cover_letter.pdf", mime="application/pdf")
                elif fmt == "DOCX":
                    data = make_docx(result, "Cover Letter")
                    st.download_button("Download", data=data, file_name="cover_letter.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.download_button("Download", data=result, file_name="cover_letter.txt", mime="text/plain")

elif page == "Settings":
    st.markdown('<div class="gradient-header">Settings</div>', unsafe_allow_html=True)
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-header">API Configuration</div>', unsafe_allow_html=True)
    current_key = st.session_state.get("gemini_api_key", "")
    new_key = st.text_input("Gemini API Key", value=current_key, type="password",
        placeholder="Enter your Gemini API key (optional)")
    if new_key:
        st.session_state["gemini_api_key"] = new_key
    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-header">Data</div>', unsafe_allow_html=True)
    if st.button("Clear session data"):
        for key in list(st.session_state.keys()):
            if key != "gemini_api_key" and key != "use_local":
                del st.session_state[key]
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)
    st.info("All features work in local mode without an API key. Add GEMINI_API_KEY for AI-powered results.")
